import os
import logging
from telegram import ReplyKeyboardMarkup, ReplyKeyboardRemove, Update
from telegram.ext import (
    Updater, CommandHandler, MessageHandler, Filters, ConversationHandler, CallbackContext
)
from docx import Document

logging.basicConfig(level=logging.INFO)

TOKEN = os.getenv("BOT_TOKEN")
TEMPLATE_PATH = "zayava_template.docx"

FIELD_KEYS = [
    "reg_number", "address",
    "rro1_1", "rro1_2", "rro1_3", "rro1_4", "rro1_5",
    "rro2_1", "rro2_2", "rro2_3", "rro2_4", "rro2_5",
    "rro3_1", "rro3_2", "rro3_3", "rro3_4", "rro3_5",
    "placement_info",
    "pay1_1", "pay1_2", "pay1_3", "pay1_4",
    "pay2_1", "pay2_2", "pay2_3", "pay2_4",
    "pay3_1", "pay3_2", "pay3_3", "pay3_4"
]

def generate_template(path="zayava_template.docx"):
    doc = Document()
    doc.add_heading("Заява на продовження ліцензії", 0)
    doc.add_paragraph("Реєстраційний номер ліцензії: {{reg_number}}")
    doc.add_paragraph("Адреса місця провадження: {{address}}")

    doc.add_paragraph("\n8. Перелік фіскальних номерів та інше:")
    table = doc.add_table(rows=4, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Фіскальні номери РРО/ПРРО"
    hdr_cells[1].text = "Дата початку"
    hdr_cells[2].text = "Дата закінчення"
    hdr_cells[3].text = "Фіскальні номери книг обліку"
    hdr_cells[4].text = "Фіскальні номери розрахункових книжок"

    for i in range(3):
        row = table.rows[i+1].cells
        row[0].text = f"{{{{rro{i+1}_1}}}}"
        row[1].text = f"{{{{rro{i+1}_2}}}}"
        row[2].text = f"{{{{rro{i+1}_3}}}}"
        row[3].text = f"{{{{rro{i+1}_4}}}}"
        row[4].text = f"{{{{rro{i+1}_5}}}}"

    doc.add_paragraph("\n9. Відомості про розташування: {{placement_info}}")

    doc.add_paragraph("\n10. Інформація про внесення платежу:")
    pay_table = doc.add_table(rows=4, cols=4)
    pay_table.rows[0].cells[0].text = "Код класифікації доходів бюджету"
    pay_table.rows[0].cells[1].text = "Сума платежу"
    pay_table.rows[0].cells[2].text = "Номер платіжної інструкції"
    pay_table.rows[0].cells[3].text = "Дата платіжної інструкції"
    for i in range(3):
        row = pay_table.rows[i+1].cells
        row[0].text = f"{{{{pay{i+1}_1}}}}"
        row[1].text = f"{{{{pay{i+1}_2}}}}"
        row[2].text = f"{{{{pay{i+1}_3}}}}"
        row[3].text = f"{{{{pay{i+1}_4}}}}"

    doc.save(path)
    print(f"Шаблон збережено як {path}")

def send_template(update, context):
    path = TEMPLATE_PATH
    if not os.path.exists(path):
        generate_template(path)
    with open(path, "rb") as f:
        update.message.reply_document(f, filename="zayava_template.docx")
    update.message.reply_text("✅ Ось шаблон для Word (docx).")

def main():
    # Створюємо шаблон якщо його нема
    if not os.path.exists(TEMPLATE_PATH):
        generate_template(TEMPLATE_PATH)

    updater = Updater(TOKEN, use_context=True)
    dp = updater.dispatcher

    dp.add_handler(CommandHandler('template', send_template))

    print("✅ Бот запущено")
    updater.start_polling()
    updater.idle()

if __name__ == "__main__":
    main()
